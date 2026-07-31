
import warnings
from pathlib import Path
from math import ceil

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.image as mpimg

from sklearn.base import clone
from sklearn.model_selection import GroupKFold, RandomizedSearchCV
from sklearn.pipeline import Pipeline
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
from sklearn.compose import TransformedTargetRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
from sklearn.ensemble import RandomForestRegressor
from sklearn.neural_network import MLPRegressor
from sklearn.inspection import permutation_importance

from xgboost import XGBRegressor
from lightgbm import LGBMRegressor
from catboost import CatBoostRegressor

import shap
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings("ignore")
plt.ioff()

# =========================================================
# GLOBAL PLOT SETTINGS
# All plot text except (a), (b), (c)... labels is increased by 2 pt
# =========================================================
BASE_FONT_SIZE = 23
TITLE_FONT_SIZE = 23
LABEL_FONT_SIZE = 23
TICK_FONT_SIZE = 23
ANNOTATION_FONT_SIZE = 23
CAPTION_FONT_SIZE = 23
SUBFIG_LABEL_FONT_SIZE = 17

plt.rcParams.update({
    "font.family": "Times New Roman",
    "font.size": BASE_FONT_SIZE,
    "axes.titlesize": TITLE_FONT_SIZE,
    "axes.labelsize": LABEL_FONT_SIZE,
    "xtick.labelsize": TICK_FONT_SIZE,
    "ytick.labelsize": TICK_FONT_SIZE,
    "legend.fontsize": BASE_FONT_SIZE,
    "figure.titlesize": TITLE_FONT_SIZE,
    "mathtext.fontset": "dejavuserif",
})

# =========================================================
# USER SETTINGS
# =========================================================
BASE_DIR = Path(__file__).resolve().parent.parent

EXCEL_PATH = BASE_DIR / "data" / "Kombinations.xlsx"
SHEET_NAME = 0
TARGET_COLS = ["f1 (Hz)", "f2 (Hz)"]

GEOMETRY_ID_COL = None
GEOMETRY_COLS = ["B (m)", "L (m)", "H (m)"]

RANDOM_STATE = 42
N_ITER = 25
N_JOBS = -1

N_OUTER_SPLITS = 5
N_INNER_SPLITS = 5

OUTPUT_DIR = BASE_DIR / "results"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PLOTS_DIR = OUTPUT_DIR / "plots"
PLOTS_DIR.mkdir(parents=True, exist_ok=True)

WORD_PATH = OUTPUT_DIR / "f1_f2_group_nested_final_report.docx"
RESULTS_XLSX = OUTPUT_DIR / "f1_f2_group_nested_final_results.xlsx"

# =========================================================
# HELPER FUNCTIONS
# =========================================================
def rmse_func(y_true, y_pred):
    return np.sqrt(mean_squared_error(y_true, y_pred))

def cc_func(y_true, y_pred):
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    if np.std(y_true) == 0 or np.std(y_pred) == 0:
        return np.nan
    return np.corrcoef(y_true, y_pred)[0, 1]

def evaluate_metrics(y_true, y_pred):
    return {
        "R2": r2_score(y_true, y_pred),
        "MAE": mean_absolute_error(y_true, y_pred),
        "RMSE": rmse_func(y_true, y_pred),
        "CC": cc_func(y_true, y_pred),
    }

def english_feature_names(feature_names):
    mapping = {
        "B (m)": "B (m)",
        "L (m)": "L (m)",
        "H (m)": "H (m)",
        "E (MPa)": "E (MPa)",
        "Ro (kg/m^3)": r"$\rho$ (kg/m$^3$)",
        "Ro (kg/m3)": r"$\rho$ (kg/m$^3$)",
    }
    return [mapping.get(f, f.replace("Ro", r"$\rho$")) for f in feature_names]

def safe_target_name(target_col):
    return target_col.replace(" ", "_").replace("(", "").replace(")", "").replace("/", "_")

def make_group_ids(dataframe):
    if GEOMETRY_ID_COL is not None:
        if GEOMETRY_ID_COL not in dataframe.columns:
            raise ValueError(f"Geometry ID column not found: {GEOMETRY_ID_COL}")
        groups = dataframe[GEOMETRY_ID_COL].astype(str)
        return groups, f"Explicit column: {GEOMETRY_ID_COL}"

    missing = [col for col in GEOMETRY_COLS if col not in dataframe.columns]
    if missing:
        raise ValueError("The following geometry columns were not found: " + ", ".join(missing))

    group_parts = []
    for col in GEOMETRY_COLS:
        values = pd.to_numeric(dataframe[col], errors="coerce").round(8)
        group_parts.append(values.astype(str))

    groups = group_parts[0]
    for part in group_parts[1:]:
        groups = groups + "_" + part
    return groups, "Combined geometry columns: " + ", ".join(GEOMETRY_COLS)

def make_parity_plot(y_true, y_pred, model_name, target_label, save_path):
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)

    lim_min = min(y_true.min(), y_pred.min())
    lim_max = max(y_true.max(), y_pred.max())
    margin = 0.05 * (lim_max - lim_min if lim_max != lim_min else 1.0)

    fig, ax = plt.subplots(figsize=(7.8, 7.0))
    ax.scatter(y_true, y_pred, alpha=0.8, s=18)
    ax.plot(
        [lim_min - margin, lim_max + margin],
        [lim_min - margin, lim_max + margin],
        linestyle="--",
        linewidth=1.8
    )

    metrics = evaluate_metrics(y_true, y_pred)
    textstr = (
        f"$R^2$ = {metrics['R2']:.4f}\n"
        f"MAE = {metrics['MAE']:.4f}\n"
        f"RMSE = {metrics['RMSE']:.4f}"
    )

    ax.text(
        0.05, 0.95, textstr,
        transform=ax.transAxes,
        verticalalignment="top",
        fontsize=ANNOTATION_FONT_SIZE,
        bbox=dict(boxstyle="round", alpha=0.15)
    )

    ax.set_title(model_name, fontsize=TITLE_FONT_SIZE)
    ax.set_xlabel(fr"Actual ${target_label}$ (Hz)", fontsize=LABEL_FONT_SIZE)
    ax.set_ylabel(fr"Predicted ${target_label}$ (Hz)", fontsize=LABEL_FONT_SIZE)
    ax.set_xlim(lim_min - margin, lim_max + margin)
    ax.set_ylim(lim_min - margin, lim_max + margin)
    ax.tick_params(axis="both", labelsize=TICK_FONT_SIZE)
    ax.grid(True, alpha=0.3)

    fig.tight_layout()
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)

def make_importance_plot(importances, feature_names_plot, title, save_path):
    idx = np.argsort(importances)[::-1]
    sorted_importances = np.asarray(importances)[idx]
    sorted_features = np.asarray(feature_names_plot)[idx]

    fig, ax = plt.subplots(figsize=(8.0, 6.2))
    ax.barh(sorted_features[::-1], sorted_importances[::-1])
    ax.set_title(title, fontsize=TITLE_FONT_SIZE)
    ax.set_xlabel("Permutation importance", fontsize=LABEL_FONT_SIZE)
    ax.set_ylabel("Features", fontsize=LABEL_FONT_SIZE)
    ax.tick_params(axis="both", labelsize=TICK_FONT_SIZE)
    ax.grid(True, axis="x", alpha=0.3)

    fig.tight_layout()
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)

def make_shap_summary_plot(model, X, title, save_path, feature_names_plot):
    if isinstance(X, pd.DataFrame):
        X_df = X.copy()
        X_df.columns = feature_names_plot
    else:
        X_df = pd.DataFrame(X, columns=feature_names_plot).copy()

    X_shap = X_df.sample(300, random_state=RANDOM_STATE) if len(X_df) > 300 else X_df.copy()

    try:
        explainer = shap.TreeExplainer(model)
        shap_values = explainer.shap_values(X_shap)
    except Exception:
        explainer = shap.Explainer(model, X_shap)
        shap_values = explainer(X_shap)

    plt.figure(figsize=(8.8, 6.6))
    shap.summary_plot(shap_values, X_shap, show=False)

    fig = plt.gcf()
    axes = fig.axes

    if len(axes) >= 1:
        ax = axes[0]
        ax.set_title(title, fontsize=TITLE_FONT_SIZE)
        ax.tick_params(axis="both", labelsize=TICK_FONT_SIZE)
        ax.xaxis.label.set_size(LABEL_FONT_SIZE)
        ax.yaxis.label.set_size(LABEL_FONT_SIZE)

    if len(axes) >= 2:
        cbar_ax = axes[-1]
        cbar_ax.tick_params(labelsize=TICK_FONT_SIZE)
        cbar_ax.yaxis.label.set_size(LABEL_FONT_SIZE)

    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close()

def combine_five_images_centered(image_paths, labels_below, output_path, fig_width=15, row_height=5.4):
    valid_items = [(p, lab) for p, lab in zip(image_paths, labels_below) if p is not None and Path(p).exists()]
    if len(valid_items) != 5:
        return None

    import matplotlib.gridspec as gridspec

    fig = plt.figure(figsize=(fig_width, row_height * 3))
    gs = gridspec.GridSpec(3, 2, figure=fig)

    axes = [
        fig.add_subplot(gs[0, 0]),
        fig.add_subplot(gs[0, 1]),
        fig.add_subplot(gs[1, 0]),
        fig.add_subplot(gs[1, 1]),
        fig.add_subplot(gs[2, :]),
    ]

    for ax, (img_path, lab) in zip(axes, valid_items):
        img = mpimg.imread(img_path)
        ax.imshow(img)
        ax.axis("off")
        ax.set_anchor("C")
        ax.text(
            0.5, -0.08, lab,
            transform=ax.transAxes,
            ha="center", va="top",
            fontsize=SUBFIG_LABEL_FONT_SIZE
        )

    plt.subplots_adjust(wspace=0.08, hspace=0.20, bottom=0.04)
    fig.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output_path

def combine_images_grid(image_paths, labels_below, output_path, ncols=2, fig_width=16, row_height=7.2):
    valid_items = [(p, lab) for p, lab in zip(image_paths, labels_below) if p is not None and Path(p).exists()]
    if not valid_items:
        return None

    n = len(valid_items)
    nrows = ceil(n / ncols)

    fig, axes = plt.subplots(nrows=nrows, ncols=ncols, figsize=(fig_width, row_height * nrows))
    axes = np.atleast_1d(axes).reshape(-1)

    for ax in axes:
        ax.axis("off")

    for (img_path, lab), ax in zip(valid_items, axes):
        img = mpimg.imread(img_path)
        ax.imshow(img)
        ax.axis("off")
        ax.text(
            0.5, -0.08, lab,
            transform=ax.transAxes,
            ha="center", va="top",
            fontsize=SUBFIG_LABEL_FONT_SIZE
        )

    plt.subplots_adjust(wspace=0.10, hspace=0.28, bottom=0.05)
    fig.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output_path

def add_heading(document, text, level=1):
    return document.add_heading(text, level=level)

def add_paragraph(document, text, bold=False):
    p = document.add_paragraph()
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p

def add_picture_with_caption(document, image_path, caption_text, width_inch=6.5):
    document.add_picture(str(image_path), width=Inches(width_inch))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption_text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

def add_dataframe_table(document, df, float_fmt="{:.4f}"):
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, (float, np.floating)):
                cells[i].text = float_fmt.format(val)
            else:
                cells[i].text = str(val)

def extract_final_pipeline(fitted_model):
    if hasattr(fitted_model, "regressor_"):
        return fitted_model.regressor_
    return fitted_model

def build_wrapped_model(estimator, need_scaling):
    steps = [("imputer", SimpleImputer(strategy="median"))]
    if need_scaling:
        steps.append(("scaler", StandardScaler()))
    steps.append(("model", clone(estimator)))
    return TransformedTargetRegressor(regressor=Pipeline(steps), transformer=None)

def mean_sd(series):
    return series.mean(), series.std(ddof=1)

# =========================================================
# READ DATA
# =========================================================
df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)
df.columns = [str(c).strip() for c in df.columns]

available_targets = [col for col in TARGET_COLS if col in df.columns]
missing_targets = [col for col in TARGET_COLS if col not in df.columns]

if missing_targets:
    print("Warning: Missing targets will be skipped:")
    for col in missing_targets:
        print(f"  - {col}")

if not available_targets:
    raise ValueError("Neither 'f1 (Hz)' nor 'f2 (Hz)' was found in the Excel file.")

for col in df.columns:
    if col != GEOMETRY_ID_COL:
        df[col] = pd.to_numeric(df[col], errors="coerce")

groups_full, group_definition = make_group_ids(df)

excluded_cols = set(TARGET_COLS)
if GEOMETRY_ID_COL is not None:
    excluded_cols.add(GEOMETRY_ID_COL)

feature_cols = [col for col in df.columns if col not in excluded_cols]
feature_names_plot = english_feature_names(feature_cols)

print(f"Group definition: {group_definition}")
print(f"Unique geometry groups: {groups_full.nunique()}")

if groups_full.nunique() < N_OUTER_SPLITS:
    raise ValueError(f"Only {groups_full.nunique()} unique groups are available, but {N_OUTER_SPLITS} outer folds were requested.")

# =========================================================
# MODELS
# =========================================================
models = {
    "RandomForest": {
        "estimator": RandomForestRegressor(random_state=RANDOM_STATE),
        "param_distributions": {
            "regressor__model__n_estimators": [100, 200, 300, 500, 800],
            "regressor__model__max_depth": [None, 3, 4, 5, 6, 8, 10, 12],
            "regressor__model__min_samples_split": [2, 3, 4, 5, 8, 10],
            "regressor__model__min_samples_leaf": [1, 2, 3, 4, 5],
            "regressor__model__max_features": ["sqrt", "log2", 0.6, 0.8, 1.0],
        },
        "tree_based": True,
        "need_scaling": False,
    },
    "XGBoost": {
        "estimator": XGBRegressor(objective="reg:squarederror", random_state=RANDOM_STATE, n_jobs=N_JOBS, verbosity=0),
        "param_distributions": {
            "regressor__model__n_estimators": [100, 200, 300, 500, 800],
            "regressor__model__max_depth": [2, 3, 4, 5, 6, 8],
            "regressor__model__learning_rate": [0.01, 0.03, 0.05, 0.1, 0.15],
            "regressor__model__subsample": [0.6, 0.7, 0.8, 0.9, 1.0],
            "regressor__model__colsample_bytree": [0.6, 0.7, 0.8, 0.9, 1.0],
            "regressor__model__min_child_weight": [1, 2, 3, 5, 7],
            "regressor__model__reg_alpha": [0.0, 0.001, 0.01, 0.1, 1.0],
            "regressor__model__reg_lambda": [0.1, 0.5, 1.0, 2.0, 5.0],
        },
        "tree_based": True,
        "need_scaling": False,
    },
    "LightGBM": {
        "estimator": LGBMRegressor(random_state=RANDOM_STATE, verbose=-1),
        "param_distributions": {
            "regressor__model__n_estimators": [100, 200, 300, 500, 800],
            "regressor__model__learning_rate": [0.01, 0.03, 0.05, 0.1, 0.15],
            "regressor__model__max_depth": [-1, 3, 4, 5, 6, 8, 10],
            "regressor__model__num_leaves": [7, 15, 31, 50, 70],
            "regressor__model__min_child_samples": [5, 10, 15, 20, 25],
            "regressor__model__subsample": [0.6, 0.7, 0.8, 0.9, 1.0],
            "regressor__model__colsample_bytree": [0.6, 0.7, 0.8, 0.9, 1.0],
            "regressor__model__reg_alpha": [0.0, 0.001, 0.01, 0.1, 1.0],
            "regressor__model__reg_lambda": [0.1, 0.5, 1.0, 2.0, 5.0],
        },
        "tree_based": True,
        "need_scaling": False,
    },
    "CatBoost": {
        "estimator": CatBoostRegressor(random_seed=RANDOM_STATE, verbose=0, allow_writing_files=False),
        "param_distributions": {
            "regressor__model__iterations": [100, 200, 300, 500, 800],
            "regressor__model__depth": [3, 4, 5, 6, 7, 8],
            "regressor__model__learning_rate": [0.01, 0.03, 0.05, 0.1, 0.15],
            "regressor__model__l2_leaf_reg": [1, 3, 5, 7, 9, 12],
            "regressor__model__subsample": [0.6, 0.7, 0.8, 0.9, 1.0],
        },
        "tree_based": True,
        "need_scaling": False,
    },
    "MLP": {
        "estimator": MLPRegressor(
            random_state=RANDOM_STATE,
            early_stopping=True,
            validation_fraction=0.15,
            n_iter_no_change=20,
            max_iter=3000,
        ),
        "param_distributions": {
            "regressor__model__hidden_layer_sizes": [(16,), (32,), (64,), (32, 16), (64, 32), (64, 32, 16)],
            "regressor__model__activation": ["relu", "tanh"],
            "regressor__model__alpha": [1e-5, 1e-4, 1e-3, 1e-2, 1e-1],
            "regressor__model__learning_rate_init": [1e-4, 5e-4, 1e-3, 5e-3, 1e-2],
            "regressor__model__solver": ["adam", "lbfgs"],
        },
        "tree_based": False,
        "need_scaling": True,
    },
}

# =========================================================
# STORAGE
# =========================================================
all_summary_tables = {}
all_fold_tables = {}
all_prediction_tables = {}
all_hyperparameter_tables = {}
all_final_hyperparameter_tables = {}
all_group_split_tables = {}
all_plot_paths = {}

# =========================================================
# TARGET LOOP
# =========================================================
for target_col in available_targets:
    target_short = target_col.split()[0]
    target_safe = safe_target_name(target_col)

    valid_mask = df[target_col].notna()
    target_df = df.loc[valid_mask].reset_index(drop=True)

    X = target_df[feature_cols].copy()
    y = target_df[target_col].copy()
    groups = groups_full.loc[valid_mask].reset_index(drop=True)

    n_unique_groups = groups.nunique()
    if n_unique_groups < N_OUTER_SPLITS:
        raise ValueError(f"{target_col}: only {n_unique_groups} unique groups are available, but {N_OUTER_SPLITS} outer folds were requested.")

    outer_cv = GroupKFold(n_splits=N_OUTER_SPLITS)
    outer_splits = list(outer_cv.split(X, y, groups))

    prediction_table = pd.DataFrame({
        f"Actual_{target_short}_Hz": y,
        "Geometry_Group": groups,
        "Outer_Fold": np.nan,
    })

    group_split_records = []

    for fold_no, (train_idx, test_idx) in enumerate(outer_splits, start=1):
        train_groups = set(groups.iloc[train_idx])
        test_groups = set(groups.iloc[test_idx])
        overlap = train_groups.intersection(test_groups)
        if overlap:
            raise RuntimeError(f"Group leakage detected in outer fold {fold_no}: {overlap}")

        prediction_table.loc[test_idx, "Outer_Fold"] = fold_no
        group_split_records.append({
            "Target": target_col,
            "Outer_Fold": fold_no,
            "Train_Rows": len(train_idx),
            "Test_Rows": len(test_idx),
            "Train_Groups": len(train_groups),
            "Test_Groups": len(test_groups),
            "Group_Overlap_Count": len(overlap),
            "Test_Group_IDs": " | ".join(sorted(map(str, test_groups))),
        })

    summary_records = []
    fold_records = []
    hyperparameter_records = []
    final_hyperparameter_records = []
    target_plot_paths = {}

    for model_name, cfg in models.items():
        print(f"{target_col} | {model_name}")
        nested_test_predictions = np.full(len(y), np.nan, dtype=float)
        fold_importances = []

        for outer_fold, (train_idx, test_idx) in enumerate(outer_splits, start=1):
            X_train = X.iloc[train_idx].copy()
            X_test = X.iloc[test_idx].copy()
            y_train = y.iloc[train_idx].copy()
            y_test = y.iloc[test_idx].copy()

            groups_train = groups.iloc[train_idx].copy()
            groups_test = groups.iloc[test_idx].copy()

            group_overlap = set(groups_train).intersection(set(groups_test))
            if group_overlap:
                raise RuntimeError(f"Group leakage detected for {model_name}, outer fold {outer_fold}: {group_overlap}")

            n_inner_groups = groups_train.nunique()
            inner_splits = min(N_INNER_SPLITS, n_inner_groups)
            if inner_splits < 2:
                raise ValueError(f"{target_col}, outer fold {outer_fold}: not enough training groups for inner GroupKFold.")

            inner_cv = GroupKFold(n_splits=inner_splits)
            wrapped_model = build_wrapped_model(cfg["estimator"], cfg["need_scaling"])

            search = RandomizedSearchCV(
                estimator=wrapped_model,
                param_distributions=cfg["param_distributions"],
                n_iter=N_ITER,
                scoring="r2",
                cv=inner_cv,
                random_state=RANDOM_STATE + outer_fold,
                n_jobs=N_JOBS,
                refit=True,
                verbose=0,
                return_train_score=False,
            )

            search.fit(X_train, y_train, groups=groups_train)
            y_pred_test = search.predict(X_test)
            nested_test_predictions[test_idx] = y_pred_test

            test_metrics = evaluate_metrics(y_test, y_pred_test)
            fold_records.append({
                "Target": target_col,
                "Model": model_name,
                "Outer_Fold": outer_fold,
                "Test_R2": test_metrics["R2"],
                "Test_MAE_Hz": test_metrics["MAE"],
                "Test_RMSE_Hz": test_metrics["RMSE"],
                "Test_CC": test_metrics["CC"],
                "Inner_Best_R2": search.best_score_,
                "Train_Size": len(train_idx),
                "Test_Size": len(test_idx),
                "Train_Group_Count": groups_train.nunique(),
                "Test_Group_Count": groups_test.nunique(),
                "Group_Overlap_Count": len(group_overlap),
            })

            hyperparameter_records.append({
                "Target": target_col,
                "Model": model_name,
                "Outer_Fold": outer_fold,
                "Inner_Best_R2": search.best_score_,
                "Best_Params": str(search.best_params_),
            })

            try:
                perm = permutation_importance(
                    search.best_estimator_,
                    X_test,
                    y_test,
                    n_repeats=20,
                    random_state=RANDOM_STATE + outer_fold,
                    scoring="r2",
                    n_jobs=N_JOBS,
                )
                fold_importances.append(perm.importances_mean)
            except Exception as exc:
                print(f"Permutation importance failed for {target_col}, {model_name}, fold {outer_fold}: {exc}")

        if np.isnan(nested_test_predictions).any():
            raise RuntimeError(f"Missing nested test predictions for {target_col} - {model_name}.")

        prediction_table[f"Pred_{model_name}"] = nested_test_predictions
        prediction_table[f"Residual_{model_name}"] = y.to_numpy() - nested_test_predictions

        model_folds = pd.DataFrame([row for row in fold_records if row["Model"] == model_name])

        test_r2_mean, test_r2_sd = mean_sd(model_folds["Test_R2"])
        test_mae_mean, test_mae_sd = mean_sd(model_folds["Test_MAE_Hz"])
        test_rmse_mean, test_rmse_sd = mean_sd(model_folds["Test_RMSE_Hz"])
        test_cc_mean, test_cc_sd = mean_sd(model_folds["Test_CC"])

        overall_test = evaluate_metrics(y, nested_test_predictions)

        summary_records.append({
            "Target": target_col,
            "Model": model_name,
            "Test_R2_Mean": test_r2_mean,
            "Test_R2_SD": test_r2_sd,
            "Test_MAE_Mean_Hz": test_mae_mean,
            "Test_MAE_SD_Hz": test_mae_sd,
            "Test_RMSE_Mean_Hz": test_rmse_mean,
            "Test_RMSE_SD_Hz": test_rmse_sd,
            "Test_CC_Mean": test_cc_mean,
            "Test_CC_SD": test_cc_sd,
            "Overall_Group_OutOfFold_Test_R2": overall_test["R2"],
            "Overall_Group_OutOfFold_Test_MAE_Hz": overall_test["MAE"],
            "Overall_Group_OutOfFold_Test_RMSE_Hz": overall_test["RMSE"],
            "Overall_Group_OutOfFold_Test_CC": overall_test["CC"],
        })

        target_plot_paths[model_name] = {}

        parity_path = PLOTS_DIR / f"{target_safe}_group_nested_parity_{model_name}.png"
        make_parity_plot(y, nested_test_predictions, model_name, target_short, parity_path)
        target_plot_paths[model_name]["parity"] = parity_path

        if fold_importances:
            importance_path = PLOTS_DIR / f"{target_safe}_group_nested_importance_{model_name}.png"
            make_importance_plot(np.mean(np.vstack(fold_importances), axis=0), feature_names_plot, model_name, importance_path)
            target_plot_paths[model_name]["importance"] = importance_path
        else:
            target_plot_paths[model_name]["importance"] = None

        final_cv_splits = min(N_INNER_SPLITS, groups.nunique())
        final_group_cv = GroupKFold(n_splits=final_cv_splits)

        final_search = RandomizedSearchCV(
            estimator=build_wrapped_model(cfg["estimator"], cfg["need_scaling"]),
            param_distributions=cfg["param_distributions"],
            n_iter=N_ITER,
            scoring="r2",
            cv=final_group_cv,
            random_state=RANDOM_STATE + 999,
            n_jobs=N_JOBS,
            refit=True,
            verbose=0,
        )
        final_search.fit(X, y, groups=groups)

        final_hyperparameter_records.append({
            "Target": target_col,
            "Model": model_name,
            "Full_Data_Group_CV_Best_R2": final_search.best_score_,
            "Final_Best_Params": str(final_search.best_params_),
        })

        if cfg["tree_based"]:
            try:
                final_pipeline = extract_final_pipeline(final_search.best_estimator_)
                imputer = final_pipeline.named_steps["imputer"]
                X_imputed = pd.DataFrame(imputer.transform(X), columns=feature_names_plot)
                final_tree_model = final_pipeline.named_steps["model"]

                shap_path = PLOTS_DIR / f"{target_safe}_shap_{model_name}_final_fit.png"
                make_shap_summary_plot(final_tree_model, X_imputed, f"SHAP Summary - {model_name}", shap_path, feature_names_plot)
                target_plot_paths[model_name]["shap"] = shap_path
            except Exception as exc:
                print(f"SHAP failed for {target_col}, {model_name}: {exc}")
                target_plot_paths[model_name]["shap"] = None
        else:
            target_plot_paths[model_name]["shap"] = None

    summary_df = pd.DataFrame(summary_records).sort_values(by="Overall_Group_OutOfFold_Test_R2", ascending=False).reset_index(drop=True)
    fold_df = pd.DataFrame(fold_records)
    hyperparameter_df = pd.DataFrame(hyperparameter_records)
    final_hyperparameter_df = pd.DataFrame(final_hyperparameter_records)
    group_split_df = pd.DataFrame(group_split_records)

    all_summary_tables[target_col] = summary_df
    all_fold_tables[target_col] = fold_df
    all_prediction_tables[target_col] = prediction_table
    all_hyperparameter_tables[target_col] = hyperparameter_df
    all_final_hyperparameter_tables[target_col] = final_hyperparameter_df
    all_group_split_tables[target_col] = group_split_df
    all_plot_paths[target_col] = target_plot_paths

    ordered_model_names = ["RandomForest", "XGBoost", "LightGBM", "CatBoost", "MLP"]

    combined_parity = PLOTS_DIR / f"{target_safe}_combined_group_nested_parity.png"
    combine_five_images_centered(
        [target_plot_paths[m]["parity"] for m in ordered_model_names],
        [f"({chr(97 + i)}) {('RF' if m == 'RandomForest' else m)}" for i, m in enumerate(ordered_model_names)],
        combined_parity
    )
    all_plot_paths[target_col]["combined_parity"] = combined_parity

    combined_importance = PLOTS_DIR / f"{target_safe}_combined_group_nested_importance.png"
    combine_five_images_centered(
        [target_plot_paths[m]["importance"] for m in ordered_model_names],
        [f"({chr(97 + i)}) {('RF' if m == 'RandomForest' else m)}" for i, m in enumerate(ordered_model_names)],
        combined_importance
    )
    all_plot_paths[target_col]["combined_importance"] = combined_importance

    tree_model_names = ["RandomForest", "XGBoost", "LightGBM", "CatBoost"]
    combined_shap = PLOTS_DIR / f"{target_safe}_combined_shap_final_fit.png"
    combine_images_grid(
        [target_plot_paths[m]["shap"] for m in tree_model_names],
        [f"({chr(97 + i)}) {('RF' if m == 'RandomForest' else m)}" for i, m in enumerate(tree_model_names)],
        combined_shap,
        ncols=2,
        fig_width=16,
        row_height=7.0
    )
    all_plot_paths[target_col]["combined_shap"] = combined_shap

# =========================================================
# PAPER-READY TABLE: TEST RESULTS ONLY
# =========================================================
paper_rows = []
for target_col, summary_df in all_summary_tables.items():
    for _, row in summary_df.iterrows():
        paper_rows.append({
            "Target": target_col,
            "Model": row["Model"],
            "Test R2 (mean ± SD)": f"{row['Test_R2_Mean']:.4f} ± {row['Test_R2_SD']:.4f}",
            "Test MAE (Hz, mean ± SD)": f"{row['Test_MAE_Mean_Hz']:.4f} ± {row['Test_MAE_SD_Hz']:.4f}",
            "Test RMSE (Hz, mean ± SD)": f"{row['Test_RMSE_Mean_Hz']:.4f} ± {row['Test_RMSE_SD_Hz']:.4f}",
            "Overall grouped out-of-fold R2": f"{row['Overall_Group_OutOfFold_Test_R2']:.4f}",
        })

paper_table_df = pd.DataFrame(paper_rows)

# =========================================================
# WRITE EXCEL
# =========================================================
with pd.ExcelWriter(RESULTS_XLSX, engine="openpyxl") as writer:
    paper_table_df.to_excel(writer, sheet_name="Paper_Ready_Table", index=False)

    for target_col in available_targets:
        short = target_col.split()[0]
        all_summary_tables[target_col].to_excel(writer, sheet_name=f"{short}_Summary", index=False)
        all_fold_tables[target_col].to_excel(writer, sheet_name=f"{short}_Fold_Metrics", index=False)
        all_prediction_tables[target_col].to_excel(writer, sheet_name=f"{short}_Predictions", index=False)
        all_hyperparameter_tables[target_col].to_excel(writer, sheet_name=f"{short}_Fold_Params", index=False)
        all_final_hyperparameter_tables[target_col].to_excel(writer, sheet_name=f"{short}_Final_Params", index=False)
        all_group_split_tables[target_col].to_excel(writer, sheet_name=f"{short}_Group_Splits", index=False)

    group_export = df.copy()
    group_export["Geometry_Group"] = groups_full
    group_export.to_excel(writer, sheet_name="Data_with_Groups", index=False)

# =========================================================
# WORD REPORT
# =========================================================
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Group-Based Nested Cross-Validation Report for f1 and f2")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

add_heading(doc, "1. Methodology", level=1)
add_paragraph(doc, f"Group definition: {group_definition}.")
add_paragraph(doc, f"The dataset contained {groups_full.nunique()} unique geometry groups.")
add_paragraph(doc, f"A group-based nested cross-validation procedure with {N_OUTER_SPLITS} outer folds and up to {N_INNER_SPLITS} inner folds was adopted.")
add_paragraph(doc, "All realizations sharing the same geometry group were kept within the same fold, so no geometry group appeared simultaneously in the training and test sets.")
add_paragraph(doc, "The main performance table reports only the outer-test results, which represent generalization to previously unseen geometry groups.")
add_paragraph(doc, "Hyperparameters were optimized only within the inner GroupKFold loop. Final hyperparameters were then re-estimated on the complete dataset using group-based cross-validation for reproducibility.")

add_heading(doc, "2. Paper-Ready Test Performance Table", level=1)
add_dataframe_table(doc, paper_table_df)

for target_col in available_targets:
    target_short = target_col.split()[0]
    add_heading(doc, f"3. Detailed Results for {target_short}", level=1)
    add_dataframe_table(doc, all_summary_tables[target_col])

    add_heading(doc, f"Outer-Fold Test Metrics for {target_short}", level=2)
    add_dataframe_table(doc, all_fold_tables[target_col])

    add_heading(doc, f"Final Hyperparameters for {target_short}", level=2)
    add_dataframe_table(doc, all_final_hyperparameter_tables[target_col])

    combined_parity = all_plot_paths[target_col]["combined_parity"]
    if combined_parity.exists():
        add_picture_with_caption(doc, combined_parity, f"Combined parity plots for {target_short} in the grouped nested cross-validation framework.", width_inch=6.5)

    combined_importance = all_plot_paths[target_col]["combined_importance"]
    if combined_importance.exists():
        add_picture_with_caption(doc, combined_importance, f"Permutation importance plots for {target_short}, evaluated on unseen geometry groups.", width_inch=6.5)

    combined_shap = all_plot_paths[target_col]["combined_shap"]
    if combined_shap.exists():
        add_picture_with_caption(doc, combined_shap, f"SHAP summary plots for the final tree-based models of {target_short}.", width_inch=6.5)

doc.save(WORD_PATH)

# =========================================================
# PRINT SUMMARY
# =========================================================
print("\n" + "=" * 90)
print("PAPER-READY GROUP-BASED TEST PERFORMANCE TABLE")
print("=" * 90)
print(paper_table_df.to_string(index=False))

print(f"\nExcel results saved to: {RESULTS_XLSX}")
print(f"Word report saved to: {WORD_PATH}")
print(f"Plots saved to: {PLOTS_DIR}")
